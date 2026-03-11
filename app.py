import streamlit as st
import pandas as pd
import json
import math
import io
import re
from datetime import datetime

# ============================================================
# 페이지 설정
# ============================================================
st.set_page_config(
    page_title="FP 자동산정 도구",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================
# 커스텀 CSS
# ============================================================
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #0D1B2A 0%, #1B3A5C 100%);
        padding: 24px 32px;
        border-radius: 16px;
        margin-bottom: 24px;
        color: white;
    }
    .main-header h1 { color: #00B4D8; margin: 0; font-size: 28px; }
    .main-header p { color: #90E0EF; margin: 4px 0 0 0; font-size: 14px; }
    .metric-card {
        background: #1E293B;
        border-radius: 12px;
        padding: 16px 20px;
        text-align: center;
        border-left: 4px solid;
    }
    .metric-card .label { font-size: 12px; color: #94A3B8; }
    .metric-card .value { font-size: 28px; font-weight: 800; }
    .metric-card .sub { font-size: 11px; color: #64748B; }
    .result-box {
        background: linear-gradient(135deg, #0D1B2A, #1B3A5C);
        border: 1px solid #00B4D8;
        border-radius: 16px;
        padding: 24px;
        color: white;
    }
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] {
        background-color: #1E293B;
        border-radius: 8px;
        color: #94A3B8;
        padding: 8px 20px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #00B4D8 !important;
        color: white !important;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================
# 상수 정의
# ============================================================
FP_WEIGHTS = {
    "ILF": {"low": 7.5, "avg": 10.0, "high": 15.0},
    "EIF": {"low": 5.0, "avg": 7.0, "high": 10.0},
    "EI":  {"low": 3.0, "avg": 4.0, "high": 6.0},
    "EO":  {"low": 4.0, "avg": 5.2, "high": 7.0},
    "EQ":  {"low": 3.0, "avg": 3.9, "high": 6.0},
}

FP_TYPE_LABELS = {
    "ILF": "내부논리파일 (Internal Logical File)",
    "EIF": "외부인터페이스파일 (External Interface File)",
    "EI":  "외부입력 (External Input)",
    "EO":  "외부출력 (External Output)",
    "EQ":  "외부조회 (External Inquiry)",
}

FP_TYPE_COLORS = {
    "ILF": "#6366f1", "EIF": "#8b5cf6", "EI": "#0891b2",
    "EO": "#059669", "EQ": "#d97706",
}

COMPLEXITY_LABELS = {"low": "낮음", "avg": "보통", "high": "높음"}

CORRECTION_OPTIONS = {
    "연계복잡성": [
        (0.88, "1. 타 기관 연계 없음"),
        (0.94, "2. 1~2개 타 기관 연계"),
        (1.00, "3. 3~5개 타 기관 연계"),
        (1.06, "4. 6~10개 타 기관 연계"),
        (1.12, "5. 10개 초과 타 기관 연계"),
    ],
    "성능 요구수준": [
        (0.91, "1. 특별한 요구사항 없음"),
        (0.95, "2. 요구사항 있으나 특별 조치 불필요"),
        (1.00, "3. 피크타임 중요, 처리시한 명시"),
        (1.05, "4. 모든 업무시간 중요, 처리시한 명시"),
        (1.09, "5. 엄격, 성능분석도구 필요"),
    ],
    "운영환경 호환성": [
        (0.94, "1. 요구사항 없음"),
        (1.00, "2. 동일 HW/SW 환경"),
        (1.06, "3. 유사 HW/SW 환경"),
        (1.13, "4. 이질적 HW/SW 환경"),
        (1.19, "5. 4 + 운영절차 문서화/모의훈련"),
    ],
    "보안성 요구수준": [
        (0.97, "1. 1가지 보안 요구"),
        (1.00, "2. 2가지 보안 요구"),
        (1.03, "3. 3가지 보안 요구"),
        (1.06, "4. 4가지 보안 요구"),
        (1.08, "5. 5가지 이상 보안 요구"),
    ],
}


def calc_scale_factor(fp):
    """SW 규모 보정계수 산출 (간이법 공식)"""
    if fp <= 0:
        return 1.0
    if fp < 500:
        return 1.28
    if fp > 3000:
        return 1.153
    return 0.4057 * (math.log(fp) - 7.1978) ** 2 + 0.8878


def extract_text_from_file(uploaded_file):
    """업로드된 파일에서 텍스트 추출"""
    name = uploaded_file.name.lower()
    content = uploaded_file.read()

    if name.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")

    elif name.endswith(".csv"):
        return content.decode("utf-8", errors="ignore")

    elif name.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        texts.append(row_text)
            return "\n".join(texts)
        except ImportError:
            st.error("python-docx 라이브러리가 필요합니다. requirements.txt에 포함되어 있습니다.")
            return ""

    elif name.endswith(".xlsx") or name.endswith(".xls"):
        try:
            all_sheets = pd.read_excel(io.BytesIO(content), sheet_name=None, header=None)
            texts = []
            for sheet_name, df in all_sheets.items():
                texts.append(f"[시트: {sheet_name}]")
                for _, row in df.iterrows():
                    vals = [str(v) for v in row.values if pd.notna(v) and str(v).strip()]
                    if vals:
                        texts.append(" | ".join(vals))
            return "\n".join(texts)
        except Exception as e:
            st.error(f"Excel 파일 읽기 오류: {e}")
            return ""

    elif name.endswith(".hwp"):
        st.warning("HWP 파일은 직접 텍스트 추출이 제한됩니다. 내용을 복사하여 텍스트로 붙여넣어 주세요.")
        return ""

    elif name.endswith(".pdf"):
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            texts = []
            for page in doc:
                texts.append(page.get_text())
            return "\n".join(texts)
        except ImportError:
            st.warning("PDF 파일 처리를 위해 PyMuPDF가 필요합니다. 텍스트를 복사하여 붙여넣어 주세요.")
            return ""

    else:
        return content.decode("utf-8", errors="ignore")


def analyze_with_claude(text, api_key):
    """Claude API로 RFP 분석하여 FP 산정"""
    import anthropic

    client = anthropic.Anthropic(api_key=api_key)

    prompt = f"""아래 RFP(제안요청서) 텍스트를 분석하여 기능점수(FP) 간이법으로 산정하기 위한 단위프로세스를 추출해주세요.

각 요구사항을 분석하여 단위프로세스를 도출하고, FP 유형(ILF, EIF, EI, EO, EQ)을 분류해주세요.

반드시 아래 JSON 형식으로만 응답하세요 (다른 텍스트 없이 JSON 배열만):

[
  {{
    "reqId": "요구사항 번호 (예: SFR-001)",
    "appName": "어플리케이션/모듈명",
    "processName": "단위프로세스명",
    "description": "단위프로세스 설명 (1~2문장)",
    "fpType": "ILF 또는 EIF 또는 EI 또는 EO 또는 EQ",
    "complexity": "low 또는 avg 또는 high",
    "rationale": "FP유형 판단 근거"
  }}
]

FP 유형 분류 기준:
- ILF: 시스템 내부에서 유지되는 데이터 그룹 (DB 테이블, 설정 저장소, 마스터 데이터 등)
- EIF: 외부 시스템에서 참조하는 데이터 (외부 API 데이터, 연계 수신 데이터 등)
- EI: 외부에서 입력되는 데이터/제어 (등록, 수정, 삭제, 업로드, 설정 변경 등)
- EO: 처리하여 출력하는 정보 (보고서, 알림, 데이터 변환, 차트 생성, 에러 처리 등)
- EQ: 단순 조회 (검색, 목록 조회, 상세 조회, 다운로드 등)

하나의 요구사항에서 여러 단위프로세스를 도출할 수 있습니다.
complexity는 대부분 "avg"로, 복잡한 것은 "high", 단순한 것은 "low"로 설정하세요.

RFP 텍스트:
{text[:20000]}"""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = message.content[0].text
    # Clean markdown fences if present
    response_text = re.sub(r'```json\s*', '', response_text)
    response_text = re.sub(r'```\s*', '', response_text).strip()

    parsed = json.loads(response_text)

    results = []
    for i, item in enumerate(parsed):
        fp_type = item.get("fpType", "EO")
        complexity = item.get("complexity", "avg")
        if fp_type not in FP_WEIGHTS:
            fp_type = "EO"
        if complexity not in ["low", "avg", "high"]:
            complexity = "avg"

        results.append({
            "No": i + 1,
            "요구사항ID": item.get("reqId", ""),
            "모듈명": item.get("appName", ""),
            "단위프로세스명": item.get("processName", ""),
            "설명": item.get("description", ""),
            "FP유형": fp_type,
            "복잡도": complexity,
            "가중치(FP)": FP_WEIGHTS[fp_type][complexity],
            "판단근거": item.get("rationale", ""),
        })

    return results


def generate_excel(df, summary, corrections, cost):
    """결과를 Excel 파일로 생성"""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        # Sheet 1: FP 산정 결과 요약
        summary_data = {
            "항목": [
                "총 기능점수", "FP당 단가", "보정계수(종합)",
                "보정 후 개발원가", f"이윤({cost['profit_rate']*100:.3f}%)",
                "SW 개발비", "직접경비", "HW 도입비",
                "공급가액(VAT제외)", "부가가치세(10%)", "총 사업비(VAT포함)"
            ],
            "금액/수치": [
                f"{summary['total_fp']:.1f} FP",
                f"{cost['fp_unit_price']:,.0f}원",
                f"{summary['total_correction']:.6f}",
                f"{summary['dev_cost']:,.0f}원",
                f"{summary['profit']:,.0f}원",
                f"{summary['sw_dev_cost']:,.0f}원",
                f"{cost['direct_expense']:,.0f}원",
                f"{cost['hw_cost']:,.0f}원",
                f"{summary['supply_cost']:,.0f}원",
                f"{summary['vat']:,.0f}원",
                f"{summary['total_cost']:,.0f}원",
            ]
        }
        pd.DataFrame(summary_data).to_excel(writer, sheet_name="FP산정결과요약", index=False)

        # Sheet 2: 보정계수
        corr_data = {
            "보정계수": ["SW규모(자동)", "연계복잡성", "성능요구", "운영환경", "보안성", "종합"],
            "값": [
                corrections["scale"], corrections["interface"],
                corrections["performance"], corrections["environment"],
                corrections["security"], summary["total_correction"]
            ]
        }
        pd.DataFrame(corr_data).to_excel(writer, sheet_name="보정계수", index=False)

        # Sheet 3: FP 유형별 집계
        type_summary = df.groupby("FP유형").agg(
            기능수=("FP유형", "count"),
            기능점수=("가중치(FP)", "sum")
        ).reset_index()
        type_summary["비중(%)"] = (type_summary["기능점수"] / type_summary["기능점수"].sum() * 100).round(1)
        type_summary.to_excel(writer, sheet_name="FP유형별집계", index=False)

        # Sheet 4: 단위프로세스 상세
        df.to_excel(writer, sheet_name="단위프로세스상세", index=False)

    output.seek(0)
    return output


# ============================================================
# 세션 상태 초기화
# ============================================================
if "requirements" not in st.session_state:
    st.session_state.requirements = []
if "rfp_text" not in st.session_state:
    st.session_state.rfp_text = ""
if "analyzed" not in st.session_state:
    st.session_state.analyzed = False


# ============================================================
# 헤더
# ============================================================
st.markdown("""
<div class="main-header">
    <h1>📊 FP 자동산정 도구</h1>
    <p>RFP 파일 업로드 → AI 자동 분석 → 기능점수(FP) 간이법 산정 → 보정계수 → 개발원가 → Excel 다운로드</p>
</div>
""", unsafe_allow_html=True)


# ============================================================
# 사이드바: API 키 & 설정
# ============================================================
with st.sidebar:
    st.markdown("### ⚙️ 설정")

    api_key = st.text_input(
        "Claude API Key",
        type="password",
        help="Anthropic API 키를 입력하세요. AI 자동 분석에 필요합니다."
    )

    st.markdown("---")
    st.markdown("### 💰 원가 파라미터")

    fp_unit_price = st.number_input(
        "FP당 단가 (원)",
        value=639102,
        step=1000,
        help="2026년 추정: 639,102원 (2025년 605,784 × 1.055)"
    )

    profit_rate = st.number_input(
        "이윤율",
        value=0.09755,
        step=0.001,
        format="%.5f",
        help="KRC 기준: 9.755%"
    )

    direct_expense = st.number_input(
        "직접경비 (원)",
        value=4242000,
        step=100000,
        help="사무실 임차료 등"
    )

    hw_cost = st.number_input(
        "HW 도입비 (원)",
        value=485000000,
        step=1000000,
        help="서버+GPU+스토리지+SAN+K8S"
    )

    st.markdown("---")
    st.markdown("### 📋 보정계수")

    corrections = {}

    for key, options in CORRECTION_OPTIONS.items():
        labels = [f"{o[1]} ({o[0]})" for o in options]
        default_idx = next((i for i, o in enumerate(options) if o[0] == 1.0), 0)
        selected = st.selectbox(key, labels, index=default_idx)
        idx = labels.index(selected)
        corrections[key] = options[idx][0]

    st.markdown("---")
    st.markdown("""
    <div style="font-size: 11px; color: #64748B;">
    <b>만든 사람:</b> 조영준<br>
    <b>산정 기준:</b> SW사업 대가산정 간이법<br>
    <b>FP당 단가:</b> 2026년 추정치<br>
    <b>보정계수:</b> SW진흥법 시행령 기준
    </div>
    """, unsafe_allow_html=True)


# ============================================================
# 메인 영역: 탭 구성
# ============================================================
tab1, tab2, tab3, tab4 = st.tabs(["📁 RFP 입력", "📋 FP 산정 결과", "💰 원가 계산", "📥 다운로드"])


# ============================================================
# TAB 1: RFP 입력
# ============================================================
with tab1:
    st.markdown("### 📁 RFP 파일 업로드 또는 텍스트 입력")

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("#### 방법 1: 파일 업로드")
        uploaded_file = st.file_uploader(
            "RFP 파일을 선택하세요",
            type=["docx", "xlsx", "xls", "pdf", "txt", "csv"],
            help="DOCX, XLSX, PDF, TXT 파일을 지원합니다. HWP는 텍스트 복사 후 방법2를 이용해주세요."
        )

        if uploaded_file:
            with st.spinner("파일에서 텍스트 추출 중..."):
                extracted = extract_text_from_file(uploaded_file)
                if extracted:
                    st.session_state.rfp_text = extracted
                    st.success(f"✅ {len(extracted):,}자 추출 완료! ({uploaded_file.name})")

    with col2:
        st.markdown("#### 방법 2: 텍스트 직접 입력")
        manual_text = st.text_area(
            "RFP 요구사항 텍스트를 붙여넣으세요",
            value=st.session_state.rfp_text,
            height=300,
            placeholder="여기에 RFP의 요구사항 부분을 복사하여 붙여넣으세요...\n\n예시:\nSFR-001: 소방 통합 데이터베이스 구축\n- 119 종합상황실 데이터 연계\n- 통합 DB 구축..."
        )
        if manual_text != st.session_state.rfp_text:
            st.session_state.rfp_text = manual_text

    st.markdown("---")

    # 분석 버튼
    col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 2])

    with col_btn1:
        if st.button("🤖 AI로 자동 분석", type="primary", use_container_width=True):
            if not api_key:
                st.error("왼쪽 사이드바에서 Claude API Key를 입력해주세요.")
            elif not st.session_state.rfp_text.strip():
                st.error("RFP 텍스트를 먼저 입력해주세요.")
            else:
                with st.spinner("🤖 Claude AI가 RFP를 분석하고 있습니다... (30초~1분 소요)"):
                    try:
                        results = analyze_with_claude(st.session_state.rfp_text, api_key)
                        st.session_state.requirements = results
                        st.session_state.analyzed = True
                        st.success(f"✅ {len(results)}개 단위프로세스 추출 완료!")
                        st.rerun()
                    except json.JSONDecodeError as e:
                        st.error(f"AI 응답 파싱 오류: {e}\n\n응답 형식이 올바르지 않습니다. 다시 시도해주세요.")
                    except Exception as e:
                        st.error(f"분석 오류: {e}")

    with col_btn2:
        if st.button("✏️ 수동으로 입력하기", use_container_width=True):
            st.session_state.requirements = []
            st.session_state.analyzed = True
            st.rerun()

    # 텍스트 미리보기
    if st.session_state.rfp_text:
        with st.expander(f"📝 추출된 텍스트 미리보기 ({len(st.session_state.rfp_text):,}자)", expanded=False):
            st.text(st.session_state.rfp_text[:3000] + ("..." if len(st.session_state.rfp_text) > 3000 else ""))


# ============================================================
# TAB 2: FP 산정 결과
# ============================================================
with tab2:
    if not st.session_state.analyzed:
        st.info("👆 'RFP 입력' 탭에서 먼저 분석을 실행해주세요.")
    else:
        reqs = st.session_state.requirements

        # FP 유형별 집계 카드
        if reqs:
            df = pd.DataFrame(reqs)
            total_fp = df["가중치(FP)"].sum()

            cols = st.columns(6)
            for i, fp_type in enumerate(["ILF", "EIF", "EI", "EO", "EQ"]):
                subset = df[df["FP유형"] == fp_type]
                with cols[i]:
                    st.metric(
                        label=f"{fp_type}",
                        value=f"{subset['가중치(FP)'].sum():.1f}",
                        delta=f"{len(subset)}건",
                    )
            with cols[5]:
                st.metric(
                    label="총 FP",
                    value=f"{total_fp:.1f}",
                    delta=f"{len(reqs)}개 프로세스",
                )

        st.markdown("---")
        st.markdown("### 📋 단위프로세스 목록 (편집 가능)")
        st.caption("각 행을 직접 수정할 수 있습니다. FP유형이나 복잡도를 변경하면 가중치가 자동 재계산됩니다.")

        # 새 프로세스 추가
        with st.expander("➕ 새 단위프로세스 추가", expanded=False):
            with st.form("add_form"):
                ac1, ac2, ac3 = st.columns(3)
                with ac1:
                    new_req_id = st.text_input("요구사항ID", placeholder="SFR-001")
                    new_module = st.text_input("모듈명", placeholder="통합 DB")
                with ac2:
                    new_process = st.text_input("단위프로세스명", placeholder="데이터 수집 처리")
                    new_desc = st.text_input("설명", placeholder="외부 데이터를 수집하여 DB에 저장")
                with ac3:
                    new_fp_type = st.selectbox("FP유형", list(FP_WEIGHTS.keys()))
                    new_complexity = st.selectbox("복잡도", list(COMPLEXITY_LABELS.keys()), format_func=lambda x: COMPLEXITY_LABELS[x])

                if st.form_submit_button("추가", type="primary"):
                    new_item = {
                        "No": len(reqs) + 1,
                        "요구사항ID": new_req_id,
                        "모듈명": new_module,
                        "단위프로세스명": new_process,
                        "설명": new_desc,
                        "FP유형": new_fp_type,
                        "복잡도": new_complexity,
                        "가중치(FP)": FP_WEIGHTS[new_fp_type][new_complexity],
                        "판단근거": "수동 입력",
                    }
                    st.session_state.requirements.append(new_item)
                    st.rerun()

        # 데이터 에디터
        if reqs:
            df = pd.DataFrame(reqs)

            edited_df = st.data_editor(
                df,
                column_config={
                    "No": st.column_config.NumberColumn("No", width=50),
                    "요구사항ID": st.column_config.TextColumn("요구사항ID", width=100),
                    "모듈명": st.column_config.TextColumn("모듈명", width=120),
                    "단위프로세스명": st.column_config.TextColumn("프로세스명", width=180),
                    "설명": st.column_config.TextColumn("설명", width=250),
                    "FP유형": st.column_config.SelectboxColumn(
                        "FP유형", options=list(FP_WEIGHTS.keys()), width=80
                    ),
                    "복잡도": st.column_config.SelectboxColumn(
                        "복잡도", options=["low", "avg", "high"], width=80
                    ),
                    "가중치(FP)": st.column_config.NumberColumn("FP", width=60, format="%.1f"),
                    "판단근거": st.column_config.TextColumn("판단근거", width=200),
                },
                num_rows="dynamic",
                use_container_width=True,
                height=500,
            )

            # 가중치 자동 재계산
            if edited_df is not None and len(edited_df) > 0:
                for idx, row in edited_df.iterrows():
                    fp_type = row.get("FP유형", "EO")
                    complexity = row.get("복잡도", "avg")
                    if fp_type in FP_WEIGHTS and complexity in FP_WEIGHTS.get(fp_type, {}):
                        edited_df.at[idx, "가중치(FP)"] = FP_WEIGHTS[fp_type][complexity]

                st.session_state.requirements = edited_df.to_dict("records")
        else:
            st.info("아직 단위프로세스가 없습니다. 위의 '추가' 버튼을 눌러 수동으로 추가하거나, RFP 입력 탭에서 AI 분석을 실행하세요.")


# ============================================================
# TAB 3: 원가 계산
# ============================================================
with tab3:
    reqs = st.session_state.requirements
    if not reqs:
        st.info("👆 'FP 산정 결과' 탭에서 먼저 단위프로세스를 추가해주세요.")
    else:
        df = pd.DataFrame(reqs)
        total_fp = df["가중치(FP)"].sum()

        # 보정계수 계산
        scale_factor = calc_scale_factor(total_fp)
        total_correction = (
            scale_factor
            * corrections["연계복잡성"]
            * corrections["성능 요구수준"]
            * corrections["운영환경 호환성"]
            * corrections["보안성 요구수준"]
        )

        # 원가 계산
        dev_cost = total_fp * fp_unit_price * total_correction
        profit = dev_cost * profit_rate
        sw_dev_cost = dev_cost + profit
        supply_cost = sw_dev_cost + direct_expense + hw_cost
        vat = supply_cost * 0.1
        total_cost = supply_cost + vat

        est_mm = dev_cost / 9089424  # 중급기술자 2026 추정

        # 보정계수 표시
        st.markdown("### ⚙️ 보정계수 현황")
        cc1, cc2, cc3, cc4, cc5, cc6 = st.columns(6)
        cc1.metric("① 규모(자동)", f"{scale_factor:.4f}", f"{total_fp:.0f}FP")
        cc2.metric("② 연계", f"{corrections['연계복잡성']}")
        cc3.metric("③ 성능", f"{corrections['성능 요구수준']}")
        cc4.metric("④ 운영환경", f"{corrections['운영환경 호환성']}")
        cc5.metric("⑤ 보안", f"{corrections['보안성 요구수준']}")
        cc6.metric("종합 보정계수", f"{total_correction:.4f}")

        st.markdown("---")

        # 원가 결과
        st.markdown("### 💰 원가 산출 결과")

        col_l, col_r = st.columns([1, 1])

        with col_l:
            st.markdown("#### SW 개발비")
            cost_data = {
                "항목": [
                    f"총 기능점수",
                    f"FP당 단가 ({fp_unit_price:,.0f}원)",
                    f"보정계수 (종합)",
                    f"보정 후 개발원가",
                    f"이윤 ({profit_rate*100:.3f}%)",
                    f"**SW 개발비**",
                ],
                "금액": [
                    f"{total_fp:.1f} FP",
                    f"{fp_unit_price:,.0f}원",
                    f"{total_correction:.6f}",
                    f"{dev_cost:,.0f}원",
                    f"{profit:,.0f}원",
                    f"**{sw_dev_cost:,.0f}원**",
                ],
            }
            st.table(pd.DataFrame(cost_data))

        with col_r:
            st.markdown("#### 사업비 총괄")
            total_data = {
                "항목": [
                    "SW 개발비",
                    "직접경비",
                    "HW 도입비",
                    "**공급가액 (VAT 제외)**",
                    "부가가치세 (10%)",
                    "**🔴 총 사업비 (VAT 포함)**",
                ],
                "금액": [
                    f"{sw_dev_cost:,.0f}원",
                    f"{direct_expense:,.0f}원",
                    f"{hw_cost:,.0f}원",
                    f"**{supply_cost:,.0f}원**",
                    f"{vat:,.0f}원",
                    f"**{total_cost:,.0f}원**",
                ],
            }
            st.table(pd.DataFrame(total_data))

        # 핵심 지표
        st.markdown("---")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("총 사업비", f"{total_cost/100000000:.1f}억원")
        m2.metric("추정 M/M", f"{est_mm:.1f}")
        m3.metric("추정 인원(10개월)", f"{est_mm/10:.1f}명")
        m4.metric("FP당 사업비", f"{total_cost/total_fp:,.0f}원/FP")


# ============================================================
# TAB 4: 다운로드
# ============================================================
with tab4:
    reqs = st.session_state.requirements
    if not reqs:
        st.info("👆 먼저 FP 산정을 완료해주세요.")
    else:
        df = pd.DataFrame(reqs)
        total_fp = df["가중치(FP)"].sum()

        scale_factor = calc_scale_factor(total_fp)
        total_correction = (
            scale_factor
            * corrections["연계복잡성"]
            * corrections["성능 요구수준"]
            * corrections["운영환경 호환성"]
            * corrections["보안성 요구수준"]
        )

        dev_cost = total_fp * fp_unit_price * total_correction
        profit = dev_cost * profit_rate
        sw_dev_cost = dev_cost + profit
        supply_cost = sw_dev_cost + direct_expense + hw_cost
        vat = supply_cost * 0.1
        total_cost = supply_cost + vat

        summary = {
            "total_fp": total_fp,
            "total_correction": total_correction,
            "dev_cost": dev_cost,
            "profit": profit,
            "sw_dev_cost": sw_dev_cost,
            "supply_cost": supply_cost,
            "vat": vat,
            "total_cost": total_cost,
        }

        corr_dict = {
            "scale": scale_factor,
            "interface": corrections["연계복잡성"],
            "performance": corrections["성능 요구수준"],
            "environment": corrections["운영환경 호환성"],
            "security": corrections["보안성 요구수준"],
        }

        cost_dict = {
            "fp_unit_price": fp_unit_price,
            "profit_rate": profit_rate,
            "direct_expense": direct_expense,
            "hw_cost": hw_cost,
        }

        st.markdown("### 📥 결과 다운로드")
        st.markdown("분석 결과를 Excel 파일로 다운로드합니다. 4개 시트로 구성됩니다.")

        st.markdown("""
        | 시트명 | 내용 |
        |--------|------|
        | FP산정결과요약 | 총 FP, 개발원가, 이윤, 총 사업비 |
        | 보정계수 | 5종 보정계수 값 |
        | FP유형별집계 | ILF/EIF/EI/EO/EQ별 건수·점수·비중 |
        | 단위프로세스상세 | 전체 단위프로세스 목록 (수정 반영) |
        """)

        excel_data = generate_excel(df, summary, corr_dict, cost_dict)

        st.download_button(
            label="📥 Excel 파일 다운로드",
            data=excel_data,
            file_name=f"FP산정결과_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            type="primary",
            use_container_width=True,
        )

        st.markdown("---")

        # CSV도 제공
        csv_data = df.to_csv(index=False).encode("utf-8-sig")
        st.download_button(
            label="📋 단위프로세스 CSV 다운로드",
            data=csv_data,
            file_name=f"FP_단위프로세스_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv",
            use_container_width=True,
        )
